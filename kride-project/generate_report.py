# -*- coding: utf-8 -*-
"""
generate_report.py
K-Ride AI개발 수행내역서 자동 생성
실행: python kride-project/generate_report.py
출력: kride-project/report/K-Ride_AI개발수행내역서.docx
"""

import os, io
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import joblib

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── 경로 ─────────────────────────────────────────────────────────────
BASE      = os.path.dirname(__file__)
DATA_PATH = os.path.join(BASE, "data", "raw_ml", "road_scored.csv")
META_PATH = os.path.join(BASE, "models", "safety_meta.pkl")
OUT_PATH  = os.path.join(BASE, "report", "K-Ride_AI개발수행내역서.docx")
CHART_DIR = os.path.join(BASE, "report", "charts")
os.makedirs(CHART_DIR, exist_ok=True)

# ── 한글 폰트 설정 ────────────────────────────────────────────────────
def set_mpl_font():
    for path in [
        "C:/Windows/Fonts/malgun.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    ]:
        if os.path.exists(path):
            fm.fontManager.addfont(path)
            prop = fm.FontProperties(fname=path)
            plt.rcParams["font.family"] = prop.get_name()
            plt.rcParams["axes.unicode_minus"] = False
            return
    plt.rcParams["axes.unicode_minus"] = False

set_mpl_font()

# ── 데이터·메타 로드 ──────────────────────────────────────────────────
df   = pd.read_csv(DATA_PATH)
meta = joblib.load(META_PATH)

# ── 차트 생성 헬퍼 ────────────────────────────────────────────────────
def save_chart(name: str, fig) -> str:
    path = os.path.join(CHART_DIR, f"{name}.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path

# 차트 1: 회귀 모델 성능 비교
def chart_model_compare():
    models = ["선형회귀", "다중회귀", "RandomForest\n(기본)", "RandomForest\n(v2 개선)"]
    r2s    = [0.0096, 0.0635, 0.1890, 0.9539]
    colors = ["#AAAAAA", "#AAAAAA", "#AAAAAA", "#4DA6FF"]
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.barh(models, r2s, color=colors)
    for bar, val in zip(bars, r2s):
        ax.text(val + 0.01, bar.get_y() + bar.get_height()/2,
                f"{val:.4f}", va="center", fontsize=9)
    ax.set_xlim(0, 1.05)
    ax.set_xlabel("R² Score")
    ax.set_title("회귀 모델 성능 비교")
    ax.axvline(0.6, color="red", lw=1, linestyle="--", label="기준선(0.6)")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return save_chart("model_compare", fig)

# 차트 2: 안전점수 분포
def chart_safety_hist():
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.hist(df["safety_score"], bins=30, color="#4DA6FF", edgecolor="white")
    ax.axvline(meta["q33"], color="#FF4444", lw=1.5, linestyle="--", label=f"q33={meta['q33']:.3f}")
    ax.axvline(meta["q66"], color="#FFA500", lw=1.5, linestyle="--", label=f"q66={meta['q66']:.3f}")
    ax.set_xlabel("안전 점수 (safety_score)")
    ax.set_ylabel("세그먼트 수")
    ax.set_title("안전 점수 분포 및 등급 임계값")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return save_chart("safety_hist", fig)

# 차트 3: 관광점수 분포
def chart_tourism_hist():
    fig, ax = plt.subplots(figsize=(6, 3))
    nz = df[df["tourism_score"] > 0]["tourism_score"]
    ax.hist(nz, bins=25, color="#57CC99", edgecolor="white")
    ax.set_xlabel("관광 점수 (tourism_score, 0 제외)")
    ax.set_ylabel("세그먼트 수")
    ax.set_title(f"관광 점수 분포 (비영 세그먼트 {len(nz):,}개)")
    fig.tight_layout()
    return save_chart("tourism_hist", fig)

# 차트 4: 안전 vs 관광 산점도
def chart_scatter():
    fig, ax = plt.subplots(figsize=(6, 4))
    sc = ax.scatter(df["safety_score"], df["tourism_score"],
                    c=df["final_score"], cmap="RdYlGn", alpha=0.4, s=8)
    plt.colorbar(sc, ax=ax, label="최종 추천 점수")
    ax.set_xlabel("안전 점수")
    ax.set_ylabel("관광 점수")
    ax.set_title("안전 점수 vs 관광 점수 (컬러: 최종 점수)")
    fig.tight_layout()
    return save_chart("scatter", fig)

# 차트 5: 분류 모델 F1 (간단 바)
def chart_f1():
    labels = ["회귀 접근\n(R²=0.19)", "분류 접근\n(F1=0.9864)"]
    vals   = [0.19, 0.9864]
    colors = ["#AAAAAA", "#4DA6FF"]
    fig, ax = plt.subplots(figsize=(5, 2.5))
    bars = ax.bar(labels, vals, color=colors)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, val + 0.01,
                f"{val:.4f}", ha="center", fontsize=10, fontweight="bold")
    ax.set_ylim(0, 1.15)
    ax.set_ylabel("성능 지표")
    ax.set_title("전략 전환 효과: 회귀 → 분류")
    fig.tight_layout()
    return save_chart("f1_compare", fig)

print("차트 생성 중...")
PATH_MODEL_COMPARE = chart_model_compare()
PATH_SAFETY_HIST   = chart_safety_hist()
PATH_TOURISM_HIST  = chart_tourism_hist()
PATH_SCATTER       = chart_scatter()
PATH_F1            = chart_f1()
print("차트 생성 완료")

# ── Word 문서 헬퍼 ────────────────────────────────────────────────────
FONT_TITLE = "휴먼명조"
FONT_BODY  = "맑은 고딕"
COLOR_RED  = RGBColor(0xC0, 0x00, 0x00)
COLOR_BLUE = RGBColor(0x00, 0x00, 0xFF)

def set_cell_border(cell, **kwargs):
    """셀 테두리 설정 헬퍼"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ["top", "left", "bottom", "right"]:
        tag = f"w:{side}"
        el = OxmlElement(tag)
        el.set(qn("w:val"),   kwargs.get("val",   "single"))
        el.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        el.set(qn("w:color"), kwargs.get("color", "000000"))
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(el)

def set_shading(cell, fill="D9D9D9"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)

def add_page_number(section):
    """하단 중앙 페이지 번호: - N -"""
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("- ")
    run.font.name = FONT_BODY; run.font.size = Pt(9)
    # 자동 페이지 번호 필드
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(fldChar1); r.append(instrText); r.append(fldChar2)
    para._p.append(r)
    run2 = para.add_run(" -")
    run2.font.name = FONT_BODY; run2.font.size = Pt(9)

def para_style(para, font_name=FONT_BODY, size=10, bold=False, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, line_spacing=1.15):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    for run in para.runs:
        run.font.name  = font_name
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), font_name)
        rPr.append(rFonts)

def add_para(doc, text, font_name=FONT_BODY, size=10, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name  = font_name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    return para

def add_heading(doc, text, level=1, color=None):
    """번호형 제목 (1. / 가. 스타일)"""
    sizes = {1: 13, 2: 11, 3: 10}
    bolds = {1: True, 2: True, 3: False}
    return add_para(doc, text,
                    size=sizes.get(level, 10),
                    bold=bolds.get(level, False),
                    color=color,
                    space_before=8 if level == 1 else 4,
                    space_after=4)

def add_red_section(doc, text):
    """빨간색 섹션 구분선 헤더"""
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.name  = FONT_BODY
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = COLOR_RED
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    rPr.append(rFonts)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(4)
    # 하단 선
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "C00000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def add_bullet(doc, text, indent_cm=0.7):
    para = doc.add_paragraph()
    run  = para.add_run(f"○ {text}")
    run.font.name  = FONT_BODY
    run.font.size  = Pt(10)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    rPr.append(rFonts)
    para.paragraph_format.left_indent   = Cm(indent_cm)
    para.paragraph_format.space_before  = Pt(1)
    para.paragraph_format.space_after   = Pt(2)
    para.paragraph_format.line_spacing  = 1.15
    return para

def add_sub_bullet(doc, text, indent_cm=1.3):
    para = doc.add_paragraph()
    run  = para.add_run(f"- {text}")
    run.font.name  = FONT_BODY
    run.font.size  = Pt(10)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    rPr.append(rFonts)
    para.paragraph_format.left_indent   = Cm(indent_cm)
    para.paragraph_format.space_before  = Pt(0)
    para.paragraph_format.space_after   = Pt(2)
    para.paragraph_format.line_spacing  = 1.15
    return para

def add_simple_table(doc, headers, rows, header_fill="4472C4", header_color="FFFFFF"):
    """헤더+데이터 테이블"""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.name = FONT_BODY
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(
            int(header_color[0:2], 16),
            int(header_color[2:4], 16),
            int(header_color[4:6], 16),
        )
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(cell, fill=header_fill)
    # 데이터
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.name = FONT_BODY
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()  # 테이블 후 여백
    return table

def add_image_in_box(doc, img_path, caption="", width_cm=14.0):
    """이미지를 테두리 박스 안에 삽입"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    if caption:
        cp = cell.add_paragraph(caption)
        cp.runs[0].font.name = FONT_BODY
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.bold = True
        cp.paragraph_format.space_after = Pt(4)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

# ══════════════════════════════════════════════════════════════════════
# 문서 생성
# ══════════════════════════════════════════════════════════════════════
doc = Document()

# ── 페이지 설정 ───────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
add_page_number(section)

# ══════════════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════════════
# 상단 여백
for _ in range(4):
    add_para(doc, "")

add_para(doc, "자전거 안전경로 추천 데이터셋 분석",
         font_name=FONT_TITLE, size=16, bold=False,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

add_para(doc, "과제AI개발 수행내역서",
         font_name=FONT_TITLE, size=28, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# 과제명/담당자 표
cover_table = doc.add_table(rows=2, cols=2)
cover_table.style = "Table Grid"
cover_data = [("과제명", "자전거도로 안전등급 분류 및 관광경로 추천 시스템 (K-Ride)"),
              ("담당자", "feed-mina")]
for ri, (label, val) in enumerate(cover_data):
    r = cover_table.rows[ri]
    r.cells[0].text = label
    r.cells[1].text = val
    for ci in range(2):
        cell = r.cells[ci]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.name = FONT_BODY
        run.font.size = Pt(11)
        if ci == 0:
            run.font.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for _ in range(2):
    add_para(doc, "")

add_para(doc, "2026년 4 월 3 일",
         font_name=FONT_TITLE, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1페이지: AI개발 수행내용
# ══════════════════════════════════════════════════════════════════════
add_red_section(doc, "AI개발 수행내용")

# 순서 박스
seq_table = doc.add_table(rows=1, cols=1)
seq_table.style = "Table Grid"
seq_cell = seq_table.rows[0].cells[0]
sp = seq_cell.add_paragraph("■ 순서")
sp.runs[0].font.name = FONT_BODY; sp.runs[0].font.size = Pt(10); sp.runs[0].font.bold = True
for i, item in enumerate(["프로젝트 개요", "데이터 분석 및 전처리 과정",
                           "모델학습 및 최적화 과정", "결과 시각화 및 평가",
                           "데이터 프로파일링 리포트"], 1):
    ip = seq_cell.add_paragraph(f"{'①②③④⑤'[i-1]} {item}")
    ip.runs[0].font.name = FONT_BODY; ip.runs[0].font.size = Pt(10)
    ip.paragraph_format.left_indent = Cm(0.5)
doc.add_paragraph()

# ── 1. 프로젝트 개요 ──────────────────────────────────────────────────
add_heading(doc, "1. 프로젝트 개요", level=1)
add_heading(doc, "  1.1 추진배경 및 목적", level=2)
for text in [
    "국내 자전거 이용 인구가 지속적으로 증가하며 자전거 교통사고 및 안전 위협 요소에 대한 사회적 관심이 높아지고 있다.",
    "기존 자전거 경로 안내 서비스는 거리·시간 최적화에 집중되어 있어, 도로 안전등급과 관광 요소를 동시에 고려한 맞춤형 추천 서비스가 부재하다.",
    "공공데이터포털 자전거도로 표준데이터, 한국관광공사 TourAPI, TAAS 교통사고 분석 데이터를 융합하여 안전 경로 추천 AI 모델을 개발하는 것을 목표로 한다.",
    "Streamlit 기반 웹 프로토타입을 구축하여 안전등급 예측 및 사용자 맞춤형(안전/관광/균형) 경로 추천 서비스를 제공한다.",
]:
    add_bullet(doc, text)

add_heading(doc, "  1.2 과제 범위", level=2)
add_simple_table(doc,
    headers=["과제구분", "내용"],
    rows=[
        ("원시 데이터 수집 및 데이터셋 구축", "자전거도로 표준데이터, 관광지 POI, 편의시설, 교통사고 데이터 수집"),
        ("데이터 전처리·표준화·상관관계 분석", "위경도 기반 Spatial Join, 결측치 처리, 피처 엔지니어링"),
        ("예측모델 선정 및 학습", "RandomForestClassifier (안전등급 분류 F1=0.9864)\nRandomForestRegressor (안전점수 회귀 R²=0.9539)"),
        ("모델 성능 평가", "R² Score·F1-macro 등 평가지표 활용"),
        ("Streamlit 웹 시스템 구축", "안전등급 예측·경로 추천·데이터 탐색 3탭 구성 및 클라우드 배포"),
    ],
)

add_heading(doc, "  1.3 과제 추진 방법", level=2)
add_heading(doc, "    1) 구축 대상 선정 기준", level=3)
for text in [
    "수도권(서울특별시·경기도) 자전거도로 세그먼트를 대상 지역으로 선정한다.",
    "전국자전거도로 표준데이터(국토교통부 공공데이터포털) 기준 서울+경기 1,647개 세그먼트를 활용한다.",
    "자전거 교통사고(TAAS)와 도로 물리적 속성(너비·길이)이 안전성에 영향을 미친다는 가설을 기반으로 피처를 구성한다.",
]:
    add_bullet(doc, text)

add_heading(doc, "    2) AI 예측 분석모델 적용 대상", level=3)
add_simple_table(doc,
    headers=["구분", "수집 데이터", "예측모델인자(독립변수)", "AI예측 분석 대상"],
    rows=[
        ("안전등급 분류", "자전거도로+TAAS 사고다발지", "width_m, length_km,\ndistrict_danger, road_attr_score", "안전등급\n(0=안전/1=보통/2=위험)"),
        ("관광 점수 산출", "TourAPI+편의시설+자전거도로", "tourist_count, cultural_count,\nleisure_count, facility_count", "tourism_score (0~1)"),
        ("경로 추천 점수", "road_scored.csv", "safety_score + tourism_score\n(모드별 가중치 조정)", "final_score → Top-10 경로 출력"),
    ],
)

add_heading(doc, "    3) AI 분석모델 구축 프로세스", level=3)
# 프로세스 플로우 표
proc_table = doc.add_table(rows=1, cols=5)
proc_table.style = "Table Grid"
steps = ["DATA\nIMPORTING", "데이터\n전처리", "모델링\n학습", "성능\n평가", "Streamlit\n배포"]
colors_proc = ["4472C4", "4472C4", "4472C4", "4472C4", "4472C4"]
for ci, (step, col) in enumerate(zip(steps, colors_proc)):
    cell = proc_table.rows[0].cells[ci]
    cell.text = step
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.font.name  = FONT_BODY; run.font.size = Pt(9); run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_shading(cell, fill=col)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 연구개발 주요 결과물
# ══════════════════════════════════════════════════════════════════════
add_red_section(doc, "연구개발 주요 결과물")

# ── 1. 데이터 수집 ────────────────────────────────────────────────────
add_heading(doc, "1. 데이터 수집", level=1)
add_heading(doc, "  가. 데이터 출처", level=2)
add_bullet(doc, "자전거도로 데이터")
add_sub_bullet(doc, "공공데이터포털, 국토교통부 전국자전거도로 표준데이터 제공")
add_bullet(doc, "관광지 POI 데이터")
add_sub_bullet(doc, "한국관광공사 국문 관광정보 서비스_GW (TourAPI)")
add_sub_bullet(doc, "서울(areaCode=1) + 경기도(areaCode=31), contentTypeId 12·14·28 수집")
add_bullet(doc, "교통사고 데이터")
add_sub_bullet(doc, "TAAS(교통사고분석시스템), 자전거 관련 교통사고 사망·중상·부상 건수")
add_bullet(doc, "편의시설 데이터")
add_sub_bullet(doc, "서울시 자전거보관함·편의시설 위치 데이터 (3,368건)")

add_heading(doc, "  나. 데이터 개요", level=2)
add_simple_table(doc,
    headers=["데이터셋", "출처", "행 수", "활용 피처"],
    rows=[
        ("road_clean.csv", "국토교통부 공공데이터포털", "5,319행", "width_m, length_km, road_type 등"),
        ("tour_poi.csv", "한국관광공사 TourAPI", "2,529건", "mapx, mapy, contentTypeId"),
        ("facility_clean.csv", "서울시 열린데이터광장", "3,368건", "x좌표, y좌표"),
        ("TAAS 사고다발지", "TAAS 교통사고분석시스템", "111건", "시군구별 사고건수·사망·중상"),
        ("road_features.csv", "Spatial Join 산출물", "1,647행", "tourist/cultural/leisure/facility_count"),
        ("road_scored.csv", "최종 모델 산출물", "1,647행", "safety_score, tourism_score, final_score"),
    ],
)

# ── 2. 데이터 전처리 및 특성 분석 ────────────────────────────────────
add_heading(doc, "2. 데이터 전처리 및 특성 분석", level=1)
add_heading(doc, "  가. 탐색적 데이터 분석 및 처리", level=2)
add_bullet(doc, "데이터셋 구성")
add_sub_bullet(doc, "자전거도로 표준데이터 전국 20,262행 중 서울특별시·경기도 필터링 → 5,319행 추출")
add_sub_bullet(doc, "위경도 결측 3,672행 제거 → 최종 1,647행 사용")
add_sub_bullet(doc, "Spatial Join: 관광지 1km 반경 / 편의시설 500m 반경 집계 (EPSG:5179 변환 후)")

add_bullet(doc, "결측치 처리")
add_simple_table(doc,
    headers=["피처", "결측 현황", "처리 방법"],
    rows=[
        ("위경도 (기점·종점)", "3,672행 결측 (69%)", "결측 행 제거 (Spatial Join 불가)"),
        ("road_type (도로종류)", "소수 결측", "'unknown'으로 대체"),
        ("강수량 (TAAS 연계용)", "비강수 기간 미입력", "0으로 대체"),
    ],
)

add_bullet(doc, "피처 통계 (road_features.csv 기준)")
add_simple_table(doc,
    headers=["피처", "평균", "최대", "설명"],
    rows=[
        ("tourist_count", "1.63", "20", "1km 반경 관광지 수"),
        ("cultural_count", "0.74", "11", "1km 반경 문화시설 수"),
        ("leisure_count", "0.10", "7", "1km 반경 레저스포츠 수"),
        ("facility_count", "1.09", "55", "500m 반경 편의시설 수"),
        ("width_m", "2.05m", "7.0m", "자전거도로 너비"),
        ("length_km", "0.83km", "36.0km", "도로 세그먼트 길이"),
    ],
)

add_heading(doc, "  나. 다중공선성 분석", level=2)
add_bullet(doc, "tourist_count·cultural_count·leisure_count 상관계수 0.7 이상 → tourism_raw로 통합 후 MinMaxScaler 정규화")
add_bullet(doc, "width_m·length_km를 road_attr_score로 결합 (너비 70% + 길이 30%)")
add_bullet(doc, "구(시군구) 단위 사고다발지 데이터를 MinMaxScaler로 정규화 → district_danger 생성 (0=안전, 1=위험)")

doc.add_page_break()

# ── 3. 데이터 학습 및 모델정의 ───────────────────────────────────────
add_heading(doc, "3. 데이터 학습 및 모델정의", level=1)
add_heading(doc, "  가. 모델 비교 및 선정", level=2)
add_bullet(doc, "회귀 모델 비교")

add_image_in_box(doc, PATH_MODEL_COMPARE, caption="[그림 1] 회귀 모델 성능 비교 (R² Score)")

add_simple_table(doc,
    headers=["분류", "모델", "독립변수", "R²", "비고"],
    rows=[
        ("회귀①", "LinearRegression", "width_m, length_km", "0.0096", "타겟과 동어반복 구조"),
        ("회귀②", "다중회귀", "관광·편의시설 피처 추가", "0.0635", "개선 미미"),
        ("회귀③", "RandomForestRegressor", "기본 피처", "0.1890", "현재 최고, 한계 확인"),
        ("회귀④", "RandomForestRegressor v2", "district_danger 반영", "0.9539", "★ 최종 채택"),
    ],
)

add_bullet(doc, "회귀 한계 원인 분석")
add_sub_bullet(doc, "safety_index = f(width_m, length_km) 공식 기반 타겟 → 입력 피처와 타겟이 동어반복 구조")
add_sub_bullet(doc, "TAAS 사고 데이터 없이 연속값 예측의 설명력이 낮음 → 3등급 분류로 전략 전환")

add_bullet(doc, "분류 전환 결과")
add_image_in_box(doc, PATH_F1, caption="[그림 2] 전략 전환 효과: 회귀 R²=0.19 → 분류 F1=0.9864")

add_heading(doc, "  나. 최종 모델 구성", level=2)
add_simple_table(doc,
    headers=["구분", "모델", "알고리즘", "성능", "저장 파일"],
    rows=[
        ("안전등급 분류", "safety_classifier", "RandomForestClassifier\n(class_weight=balanced)", "F1-macro=0.9864", "safety_classifier.pkl"),
        ("안전점수 회귀", "safety_regressor", "RandomForestRegressor", "R²=0.9539", "safety_regressor.pkl"),
        ("피처 스케일러", "safety_scaler", "MinMaxScaler", "-", "safety_scaler.pkl"),
        ("관광점수 스케일러", "tourism_scaler", "MinMaxScaler", "-", "tourism_scaler.pkl"),
    ],
)

add_heading(doc, "  다. 하이퍼파라미터 설정", level=2)
add_simple_table(doc,
    headers=["하이퍼파라미터", "설정값", "선정 근거"],
    rows=[
        ("n_estimators", "200", "과적합-성능 균형"),
        ("max_depth", "10", "적절한 복잡도 제한"),
        ("min_samples_leaf", "3", "리프 노드 최소 샘플 보장"),
        ("class_weight", "balanced", "3등급 클래스 불균형 대응"),
        ("random_state", "42", "재현성 확보"),
    ],
)

add_heading(doc, "  라. 안전등급 구간화 기준", level=2)
add_bullet(doc, f"3분위(33%/66%) 기준 자동 구간화: q33={meta['q33']:.4f}, q66={meta['q66']:.4f}")
add_simple_table(doc,
    headers=["등급", "기준", "설명"],
    rows=[
        ("0 (안전)", f"safety_score ≥ {meta['q66']:.4f}", "상위 34% 세그먼트"),
        ("1 (보통)", f"{meta['q33']:.4f} ≤ safety_score < {meta['q66']:.4f}", "중간 33% 세그먼트"),
        ("2 (위험)", f"safety_score < {meta['q33']:.4f}", "하위 33% 세그먼트"),
    ],
)

doc.add_page_break()

# ── 4. 최종 결과 시각화 ───────────────────────────────────────────────
add_heading(doc, "4. 최종 결과 시각화", level=1)
add_heading(doc, "  가. 안전점수 분포", level=2)
add_image_in_box(doc, PATH_SAFETY_HIST, caption="[그림 3] 안전점수 히스토그램 및 등급 임계값")

add_heading(doc, "  나. 관광점수 분포", level=2)
add_image_in_box(doc, PATH_TOURISM_HIST, caption="[그림 4] 관광점수 분포 (비영 세그먼트 기준)")

add_heading(doc, "  다. 안전점수 vs 관광점수 분포", level=2)
add_image_in_box(doc, PATH_SCATTER, caption="[그림 5] 안전점수 vs 관광점수 산점도 (컬러: 최종 추천점수)")

add_heading(doc, "  라. Streamlit 웹 애플리케이션", level=2)
add_bullet(doc, "배포 URL: https://humaneducation-kride.streamlit.app/")
add_simple_table(doc,
    headers=["탭", "기능", "사용 모델"],
    rows=[
        ("탭 1 – 안전등급 예측", "도로 너비·길이·위험도 슬라이더 입력\n→ 🟢안전 / 🟡보통 / 🔴위험 등급 출력 + 신뢰도 확률 바", "safety_classifier.pkl\nsafety_scaler.pkl"),
        ("탭 2 – 경로 추천 Top-10", "모드(안전/균형/관광) 선택\n→ 가중치 재계산 후 상위 10개 경로 테이블 + 분포 차트", "road_scored.csv"),
        ("탭 3 – 데이터 탐색", "히스토그램·산점도·기술통계 시각화", "road_scored.csv"),
    ],
)

add_heading(doc, "  마. 모드별 가중치", level=2)
add_simple_table(doc,
    headers=["모드", "안전 점수 가중치", "관광 점수 가중치", "대상 사용자"],
    rows=[
        ("🛡️ 안전 우선", "70%", "30%", "가족·어린이·초보 라이더"),
        ("⚖️ 균형", "50%", "50%", "일반 레저 라이더"),
        ("🗺️ 관광 우선", "30%", "70%", "관광 목적 라이더"),
    ],
)

doc.add_page_break()

# ── 5. 데이터 프로파일링 리포트 ──────────────────────────────────────
add_heading(doc, "5. 데이터 프로파일링 리포트", level=1)
add_bullet(doc, "자전거도로 데이터 (road_scored.csv)")
add_simple_table(doc,
    headers=["항목", "값"],
    rows=[
        ("전체 행 수", "1,647행"),
        ("최종 컬럼 수", "19개"),
        ("대상 지역", "서울특별시 + 경기도"),
        ("safety_score 평균", f"{df['safety_score'].mean():.3f}"),
        ("safety_score 범위", f"{df['safety_score'].min():.3f} ~ {df['safety_score'].max():.3f}"),
        ("tourism_score 평균", f"{df['tourism_score'].mean():.3f}"),
        ("tourism_score = 0 비율", f"{(df['tourism_score']==0).sum()/len(df)*100:.1f}% ({(df['tourism_score']==0).sum():,}개)"),
        ("final_score 평균", f"{df['final_score'].mean():.3f}"),
        ("final_score 최대", f"{df['final_score'].max():.3f}"),
    ],
)

add_bullet(doc, "주요 변수 기술통계")
stat_cols = ["width_m", "length_km", "tourist_count", "cultural_count",
             "facility_count", "safety_score", "tourism_score", "final_score"]
stat_df = df[stat_cols].describe().T.round(3)
stat_rows = []
for col in stat_cols:
    row = stat_df.loc[col]
    stat_rows.append((col,
                      f"{row['count']:.0f}",
                      f"{row['mean']:.3f}",
                      f"{row['std']:.3f}",
                      f"{row['min']:.3f}",
                      f"{row['50%']:.3f}",
                      f"{row['max']:.3f}"))
add_simple_table(doc,
    headers=["피처", "count", "mean", "std", "min", "median", "max"],
    rows=stat_rows,
)

doc.add_page_break()

# ── 6. 결론 및 향후 계획 ─────────────────────────────────────────────
add_heading(doc, "6. 결론 및 향후 계획", level=1)
add_heading(doc, "  가. 모델 개발 결과 요약", level=2)
add_bullet(doc, "회귀 성능 한계(R²=0.19) 확인 후 3등급 분류 전략으로 전환하여 F1-macro=0.9864를 달성하였다.")
add_bullet(doc, "TAAS 교통사고 데이터(구 단위 위험도)를 district_danger 피처로 반영, R²=0.9539의 회귀 성능 개선을 확인하였다.")
add_bullet(doc, "관광지 1km 반경·편의시설 500m 반경 Spatial Join을 통해 관광 점수(tourism_score)를 산출하였다.")
add_bullet(doc, "Streamlit 3탭 앱을 구축하고 클라우드 배포(humaneducation-kride.streamlit.app)를 완료하였다.")

add_heading(doc, "  나. 실무 적용 방안", level=2)
add_bullet(doc, "수도권 자전거 라이더 대상으로 사용자 목적(안전/관광/균형)에 따른 맞춤형 경로 추천 서비스로 활용 가능하다.")
add_bullet(doc, "따릉이·공공 자전거 서비스와 연계하여 경로 추천 알고리즘의 백엔드 엔진으로 통합할 수 있다.")
add_bullet(doc, "공공데이터포털의 최신 자전거도로 데이터로 주기적 업데이트 시 모델 성능을 유지할 수 있다.")

add_heading(doc, "  다. 한계점 및 개선 방향", level=2)
add_sub_bullet(doc, "TAAS Spatial Join(100m 반경 사고 집계)이 미완성으로, 노선 단위 세밀한 위험도 반영이 제한적이다. 향후 사고 좌표 정제 후 연계 예정이다.")
add_sub_bullet(doc, "경기도 자전거 교통사고 데이터가 미확보 상태로, 현재 서울 25구 기준 위험도만 반영하고 나머지는 도로 속성만 사용한다.")
add_sub_bullet(doc, "YOLOv8 객체 탐지 기반 실시간 위험 감지(보행자·장애물) 기능은 2주차 과제로 이월하였다.")
add_sub_bullet(doc, "실시간 혼잡도 데이터(TOPIS, 따릉이 이용 통계)를 추가하면 시간대별 경로 추천 정확도를 높일 수 있다.")

# ══════════════════════════════════════════════════════════════════════
# 저장
# ══════════════════════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"\n[완료] 보고서 생성 완료 -> {OUT_PATH}")
